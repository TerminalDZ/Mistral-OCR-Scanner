# app/main.py
import os
import json
import asyncio
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from fastapi import FastAPI, File, UploadFile, HTTPException, Query
from fastapi.responses import FileResponse, HTMLResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
from dotenv import load_dotenv

from .utils import save_upload_file, result_path, init_db, insert_job, list_jobs, get_job, DB_PATH
from .tasks import process_document
from .mistral_client import client

from docx import Document

load_dotenv()

app = FastAPI(title="Mistral OCR Scanner — Final")

# Serve frontend static files from ../frontend
FRONTEND_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "frontend")
app.mount("/static", StaticFiles(directory=os.path.join(FRONTEND_DIR)), name="static")

# Allow CORS for development (same-origin not required)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# DB init (will create table and run migrations if needed)
init_db()

# executor will be created at startup
executor = None

@app.on_event("startup")
async def startup_event():
    global executor
    if executor is None:
        maxw = int(os.environ.get("MAX_WORKERS", "3"))
        executor = ThreadPoolExecutor(max_workers=maxw)

@app.on_event("shutdown")
async def shutdown_event():
    global executor
    if executor:
        try:
            executor.shutdown(wait=False)
        except Exception:
            pass

@app.get("/", response_class=HTMLResponse)
async def index():
    index_path = os.path.join(FRONTEND_DIR, "index.html")
    if os.path.exists(index_path):
        return FileResponse(index_path, media_type="text/html")
    return HTMLResponse("<h3>Frontend not found. Put frontend/index.html in project.</h3>", status_code=404)

@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...), do_annotations: bool = False, do_qna: bool = False, annotation_schema: str = None):
    fp, job_id, original_name = save_upload_file(file)
    size_bytes = os.path.getsize(fp)
    # insert job with empty title for now (will be set after processing)
    insert_job(job_id, None, original_name, fp, size_bytes)

    annotation_schema_obj = None
    if annotation_schema:
        try:
            annotation_schema_obj = json.loads(annotation_schema)
        except Exception:
            raise HTTPException(status_code=400, detail="annotation_schema must be valid JSON")

    loop = asyncio.get_running_loop()
    # schedule processing in thread pool; non-blocking
    loop.run_in_executor(executor, process_document, fp, None, job_id, do_annotations, annotation_schema_obj, do_qna)
    return {"job_id": job_id}

@app.get("/api/status/{job_id}")
async def job_status(job_id: str):
    path = result_path(job_id)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            try:
                data = json.load(f)
            except Exception:
                data = {}
        status = data.get("status", "completed") if isinstance(data, dict) else "completed"
        return {"job_id": job_id, "status": status, "result_url": f"/api/result/{job_id}"}
    return {"job_id": job_id, "status": "pending", "result_url": None}

@app.get("/api/result/{job_id}")
async def get_result(job_id: str):
    path = result_path(job_id)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Result not ready")
    # Return JSON file for debugging; frontend won't show raw JSON by default
    return FileResponse(path, media_type="application/json", filename=f"{job_id}.json")

@app.get("/api/jobs")
async def api_jobs(limit: int = 100):
    rows = list_jobs(limit=limit)
    return {"jobs": rows, "db": str(DB_PATH)}

@app.post("/api/qna")
async def qna(body: dict):
    job_id = body.get("job_id")
    question = body.get("question")
    if not job_id or not question:
        raise HTTPException(status_code=400, detail="job_id and question required")

    path = result_path(job_id)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Job not found")

    with open(path, "r", encoding="utf-8") as f:
        try:
            doc = json.load(f)
        except Exception:
            doc = {}

    doc_url = doc.get("document_url")
    if not doc_url:
        job = get_job(job_id)
        if not job:
            raise HTTPException(status_code=400, detail="No job record found and no document_url in result")
        filepath = job.get("filepath")
        if not filepath or not os.path.exists(filepath):
            raise HTTPException(status_code=400, detail="Document URL not available for QnA — run job with do_qna=True or upload as URL")
        try:
            up = client.files.upload(file={"file_name": os.path.basename(filepath), "content": open(filepath, "rb")}, purpose="ocr")
            signed_obj = client.files.get_signed_url(file_id=up.id)
            doc_url = signed_obj.url
            doc["document_url"] = doc_url
            with open(path, "w", encoding="utf-8") as f:
                json.dump(doc, f, ensure_ascii=False, indent=2)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to upload file for QnA: {e}")

    try:
        messages = [{"role":"user","content":[{"type":"text","text": question},{"type":"document_url","document_url":doc_url}]}]
        chat_response = client.chat.complete(model=os.environ.get("DOC_QNA_MODEL","mistral-small-latest"), messages=messages)
        answer = chat_response.choices[0].message.content
        # persist qna
        try:
            doc.setdefault("qna_history", []).append({"question": question, "answer": answer})
            with open(path, "w", encoding="utf-8") as f:
                json.dump(doc, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        return {"answer": answer}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"QnA failed: {e}")

@app.get("/api/download/{job_id}")
async def api_download(job_id: str, format: str = Query("md", regex="^(md|docx)$")):
    path = result_path(job_id)
    if not os.path.exists(path):
        raise HTTPException(status_code=404, detail="Result not ready")
    with open(path, "r", encoding="utf-8") as f:
        doc = json.load(f)
    # prefer stored full_markdown
    markdown = doc.get("full_markdown") or doc.get("qna_summary") or None
    if not markdown:
        # try to extract from result['ocr'] crudely
        markdown = doc.get("title", "") + "\n\n" + (doc.get("qna_summary") or "")
    if format == "md":
        data = markdown.encode("utf-8")
        return StreamingResponse(BytesIO(data), media_type="text/markdown", headers={"Content-Disposition": f"attachment; filename={job_id}.md"})
    else:
        # build simple docx
        document = Document()
        # Title
        title = doc.get("title") or ""
        if title:
            p = document.add_paragraph()
            run = p.add_run(title)
            run.bold = True
        # naive markdown -> paragraphs conversion
        for line in markdown.splitlines():
            line = line.rstrip()
            if line.startswith("#"):
                # heading level
                level = len(line.split(" ")[0].strip("#"))
                text = line[level+1:].strip() if len(line.split(" "))>1 else line.lstrip("#").strip()
                p = document.add_paragraph()
                r = p.add_run(text)
                r.bold = True
            else:
                document.add_paragraph(line)
        bio = BytesIO()
        document.save(bio)
        bio.seek(0)
        return StreamingResponse(bio, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": f"attachment; filename={job_id}.docx"})
